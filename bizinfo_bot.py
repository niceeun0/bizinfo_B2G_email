#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
bizinfo_bot.py
=========================================================
기업마당(BizInfo) 신규 지원사업 공고 자동 수집 -> 첨부파일 원문 추출
-> AI 요약(OpenRouter) -> HTML 뉴스레터 생성 -> Gmail SMTP 발송

GitHub Actions(Ubuntu-latest) 환경에서 실행되도록 설계되었습니다.

[중요 참고사항]
1. "Temporary failure in name resolution" / "Connection timed out" 에러는
   대부분 GitHub Actions 러너의 일시적 로컬 DNS 조회 실패이며, 서버가
   실제로 클라우드 IP 대역 자체를 차단하는 경우도 있을 수 있습니다.
   완전한 우회를 보장할 수는 없으나, 아래 방식으로 성공률을 크게
   높였습니다:
     - Google/Cloudflare DoH(DNS-over-HTTPS)로 직접 IP를 조회
     - curl --resolve 로 로컬 DNS를 건너뛰고 강제로 해당 IP에 연결
     - curl 실패 시 requests 로 재시도 (다른 TLS/HTTP 스택 사용)
     - 지수 백오프 재시도(최대 5회)
   그래도 실패한다면 서버 측에서 클라우드 IP 대역 자체를 차단하고
   있는 것이므로, 국내 리전 프록시/자체 서버 경유가 필요합니다.

2. 기업마당 API의 실제 JSON 필드명은 시점에 따라 다를 수 있습니다.
   아래 코드는 흔히 쓰이는 필드명(pblancId, pblancNm, jrsdInsttNm,
   bsnsSumryCn, reqstBeginEndDe, creatPnttm, pubDate, printFlpthNm,
   printFileNm, rceptEngnHmpgUrl 등)을 우선 시도하고, 없으면 유사한
   키를 탐색하도록 방어적으로 작성했습니다. 실제 응답 구조가 다르면
   FIELD_CANDIDATES 부분만 수정하면 됩니다.

3. 환경변수 GEMINI_API_KEY 에는 (요청하신 대로 변수명은 유지하되)
   실제로는 OpenRouter API 키를 넣어서 사용합니다.
=========================================================
"""

import os
import re
import io
import csv
import sys
import json
import time
import base64
import socket
import difflib
import zipfile
import smtplib
import subprocess
import traceback
from datetime import datetime, timedelta, timezone
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from urllib.parse import quote as urlquote

import requests

# ---- 선택적 문서 파싱 라이브러리 (없으면 해당 포맷만 건너뜀) ----
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx  # python-docx
except ImportError:
    docx = None

try:
    import olefile
except ImportError:
    olefile = None

try:
    import pytesseract
    from pdf2image import convert_from_path
except ImportError:
    pytesseract = None
    convert_from_path = None

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
except ImportError:
    openpyxl = None


# =========================================================
# 0. 설정 / 환경변수
# =========================================================
BIZINFO_HOST = "www.bizinfo.go.kr"
BIZINFO_API_URL = (
    "https://www.bizinfo.go.kr/uss/rss/bizinfoApi.do"
    "?crtfcKey=4vc2gy&dataType=json&searchCnt=100"
)

MAIL_USER = os.environ.get("MAIL_USER", "")
EMAIL_PASS = os.environ.get("EMAIL_PASS", "")
MAIL_RECEIVER = os.environ.get("MAIL_RECEIVER", "")

# 요청사항: OpenRouter를 쓰되 환경변수 이름은 GEMINI_API_KEY 그대로 사용
OPENROUTER_API_KEY = os.environ.get("GEMINI_API_KEY", "")
# openrouter/free 는 OpenRouter가 공식 지원하는 "Free Models Router" 슬러그로,
# 요청 특성에 맞는 무료 모델을 자동으로 골라 라우팅합니다.
# (특정 모델을 고정하고 싶다면 "모델명:free" 형태로 환경변수 override 가능,
#  단 개별 무료 모델 슬러그는 자주 폐기/변경되므로 openrouter/free 권장)
OPENROUTER_MODEL = os.environ.get("OPENROUTER_MODEL", "openrouter/free")
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

REQUEST_TIMEOUT = 20
# 같은 GitHub Actions Job 안에서는 출발지 IP가 바뀌지 않으므로, 서버가
# 그 IP 대역 자체를 차단한 경우 재시도를 아무리 반복해도 결과가 같습니다.
# (KR_PROXY_URL이 설정되어 있으면 그쪽을 우선 시도하므로 이 값과 무관하게
# 프록시 경로가 성공할 수 있습니다.) 예전엔 5회였는데, 실패가 반복 확인된
# 뒤로는 CI 실행 시간 낭비를 줄이기 위해 기본값을 낮췄습니다.
MAX_RETRY = int(os.environ.get("MAX_RETRY", "3"))
BACKOFF_BASE = 3  # seconds

# 국내(한국) 소재 프록시 서버 주소 (예: http://user:pass@1.2.3.4:8080).
# 정부기관 사이트가 해외/클라우드 IP 대역 자체를 차단하는 경우, DNS 트릭으로는
# 절대 해결되지 않고 실제로 한국 IP를 통해 나가는 것만이 유일한 해결책입니다.
# 설정해두면 curl/requests 모두 이 프록시를 통해 나갑니다.
KR_PROXY_URL = os.environ.get("KR_PROXY_URL", "").strip()

WORKDIR = os.path.abspath(os.path.dirname(__file__))
TMP_DIR = os.path.join(WORKDIR, "_tmp_attachments")
os.makedirs(TMP_DIR, exist_ok=True)


def log(msg):
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] {msg}", flush=True)


def strip_html(text):
    """기업마당 API의 일부 텍스트 필드(예: reqstMthPapersCn)에 <br>, <p> 등
    HTML 태그가 섞여 오는 경우가 있어 정리합니다."""
    if not text:
        return ""
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"&nbsp;", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# =========================================================
# 1. 네트워크 우회 / 견고한 API 호출
# =========================================================
def resolve_ip_via_doh(hostname):
    """Google / Cloudflare DoH를 이용해 로컬 DNS를 건너뛰고 IP를 직접 조회."""
    doh_endpoints = [
        f"https://dns.google/resolve?name={hostname}&type=A",
        f"https://cloudflare-dns.com/dns-query?name={hostname}&type=A",
    ]
    for url in doh_endpoints:
        try:
            headers = {"accept": "application/dns-json"}
            resp = requests.get(url, headers=headers, timeout=10)
            if resp.status_code == 200:
                data = resp.json()
                answers = data.get("Answer", [])
                for ans in answers:
                    ip = ans.get("data")
                    # A 레코드(IPv4)만 사용
                    if ip and re.match(r"^\d{1,3}(\.\d{1,3}){3}$", ip):
                        log(f"DoH로 IP 확인 성공: {hostname} -> {ip} (via {url})")
                        return ip
        except Exception as e:
            log(f"DoH 조회 실패 ({url}): {e}")
            continue
    return None


def curl_get(url, resolved_ip=None, timeout=REQUEST_TIMEOUT, use_proxy=False):
    """subprocess로 curl 호출. resolved_ip가 있으면 --resolve로 강제 연결."""
    cmd = [
        "curl",
        "-sS",
        "-L",
        "--max-time", str(timeout),
        "--connect-timeout", str(min(timeout, 10)),
        "-A", "Mozilla/5.0 (Windows NT 10.0; Win64; x64) BizInfoBot/1.0",
    ]
    if use_proxy and KR_PROXY_URL:
        cmd += ["--proxy", KR_PROXY_URL]
    if resolved_ip and not use_proxy:
        # 프록시를 쓸 때는 목적지 IP를 프록시 서버가 직접 조회하므로 --resolve 불필요
        cmd += ["--resolve", f"{BIZINFO_HOST}:443:{resolved_ip}"]
    cmd += [url]

    result = subprocess.run(cmd, capture_output=True, timeout=timeout + 5)
    if result.returncode != 0:
        raise RuntimeError(
            f"curl 실패 (returncode={result.returncode}): "
            f"{result.stderr.decode('utf-8', errors='ignore')[:300]}"
        )
    return result.stdout


def fetch_with_retry(url):
    """
    다단계 견고한 호출:
      0) (KR_PROXY_URL이 설정된 경우) 국내 프록시를 통한 curl 호출 - 최우선
      1) curl + DoH로 조회한 IP를 --resolve로 강제 사용
      2) curl (일반 시스템 DNS)
      3) requests (일반 시스템 DNS, 다른 네트워크 스택)
    각 단계 실패 시 지수 백오프로 재시도.

    주의: 1~3번은 모두 "DNS를 못 찾는" 문제에는 도움이 되지만, 서버(또는
    중간 방화벽)가 IP 대역 자체를 차단해 TCP 연결 자체가 타임아웃 나는
    경우에는 전혀 도움이 되지 않습니다. 이 경우 KR_PROXY_URL로 실제 한국
    IP를 통해 나가는 것만이 유일한 해결책입니다.
    """
    resolved_ip = resolve_ip_via_doh(BIZINFO_HOST)

    last_err = None
    for attempt in range(1, MAX_RETRY + 1):
        # 0) 국내 프록시 경유 (설정된 경우 최우선 시도)
        if KR_PROXY_URL:
            try:
                log(f"[시도 {attempt}] KR_PROXY_URL 프록시 경유 curl 호출")
                raw = curl_get(url, use_proxy=True)
                return raw
            except Exception as e:
                last_err = e
                log(f"프록시 경유 curl 실패: {e}")

        # 1) curl + --resolve (DoH IP) - 로컬 DNS 실패에만 효과 있음
        if resolved_ip:
            try:
                log(f"[시도 {attempt}] curl --resolve 방식으로 호출")
                raw = curl_get(url, resolved_ip=resolved_ip)
                return raw
            except Exception as e:
                last_err = e
                log(f"curl --resolve 실패: {e}")

        # 2) curl 일반 호출
        try:
            log(f"[시도 {attempt}] curl 일반 호출")
            raw = curl_get(url, resolved_ip=None)
            return raw
        except Exception as e:
            last_err = e
            log(f"curl 일반 호출 실패: {e}")

        # 3) requests 폴백
        try:
            log(f"[시도 {attempt}] requests 폴백 호출")
            proxies = (
                {"https": KR_PROXY_URL, "http": KR_PROXY_URL} if KR_PROXY_URL else None
            )
            resp = requests.get(
                url,
                timeout=REQUEST_TIMEOUT,
                headers={"User-Agent": "Mozilla/5.0 BizInfoBot/1.0"},
                proxies=proxies,
            )
            resp.raise_for_status()
            return resp.content
        except Exception as e:
            last_err = e
            log(f"requests 폴백 실패: {e}")

        sleep_time = BACKOFF_BASE * attempt
        log(f"{sleep_time}초 대기 후 재시도합니다...")
        time.sleep(sleep_time)

    if not KR_PROXY_URL:
        log(
            "[진단] 모든 방식이 '연결 타임아웃'(연결 거부가 아님)으로 실패했습니다. "
            "이는 DNS 문제가 아니라 서버 또는 중간 방화벽이 이 요청의 출발지 IP "
            "대역 자체를 막고 있다는 뜻입니다. 같은 Job 안에서는 출발지 IP가 "
            "바뀌지 않으므로 재시도를 더 반복해도 결과는 같습니다. "
            "이 실패가 반복된다면(한 번의 우연한 타이밍이 아니라면), "
            "KR_PROXY_URL 환경변수(시크릿)에 국내(한국) 소재 프록시 서버 주소를 "
            "등록해야 합니다. 아직 등록하지 않았다면 지금이 등록할 시점입니다."
        )

    raise RuntimeError(f"모든 재시도 실패. 마지막 에러: {last_err}")


def fetch_bizinfo_items():
    raw = fetch_with_retry(BIZINFO_API_URL)
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        raise RuntimeError(f"JSON 파싱 실패: {e} / 원본 앞부분: {raw[:200]}")

    # 응답 구조 방어적 처리: jsonArray 혹은 최상위 리스트 등
    if isinstance(data, dict):
        items = data.get("jsonArray") or data.get("items") or data.get("result") or []
    elif isinstance(data, list):
        items = data
    else:
        items = []

    log(f"API에서 총 {len(items)}건의 공고를 수신했습니다.")
    return items


# =========================================================
# 2. 날짜 필터링 (메일 발송 전날 등록된 신규 공고만)
# =========================================================
DATE_FIELD_CANDIDATES = ["pubDate", "creatPnttm", "regDt", "creatPnttm1", "rceptDt"]


def get_item_date_str(item):
    for key in DATE_FIELD_CANDIDATES:
        if item.get(key):
            return str(item[key])
    return None


def parse_flexible_date(date_str):
    """다양한 형식(YYYYMMDD, YYYY-MM-DD, RFC1123 등)을 안전하게 파싱."""
    if not date_str:
        return None
    date_str = date_str.strip()

    # RFC 1123 (RSS pubDate 형식) 예: "Wed, 29 Jul 2026 10:00:00 +0900"
    try:
        from email.utils import parsedate_to_datetime
        dt = parsedate_to_datetime(date_str)
        if dt:
            return dt.date()
    except Exception:
        pass

    candidates = [
        ("%Y-%m-%d %H:%M:%S", 19),
        ("%Y-%m-%d", 10),
        ("%Y.%m.%d", 10),
        ("%Y/%m/%d", 10),
        ("%Y%m%d", 8),
    ]
    for fmt, length in candidates:
        try:
            return datetime.strptime(date_str[:length], fmt).date()
        except Exception:
            continue

    # 마지막 시도: 숫자만 추출해서 YYYYMMDD로 처리
    digits = re.sub(r"\D", "", date_str)
    if len(digits) >= 8:
        try:
            return datetime.strptime(digits[:8], "%Y%m%d").date()
        except Exception:
            pass
    return None


def filter_new_items(items, target_date):
    """target_date(등록일)에 등록된 공고만 필터링."""
    filtered = []
    for item in items:
        date_str = get_item_date_str(item)
        item_date = parse_flexible_date(date_str)
        if item_date == target_date:
            filtered.append(item)
    log(f"{target_date} 등록 공고 {len(filtered)}건 필터링 완료.")
    return filtered


# =========================================================
# 3. 첨부파일 URL 추출
# =========================================================
FILEBLANK_RE = re.compile(
    r"fileBlank\(\s*'([^']+)'\s*\+\s*'([^']*)'\s*\+\s*'([^']+)'\s*,\s*'([^']+)'\s*\)"
)
FILEBLANK_SIMPLE_RE = re.compile(
    r"fileBlank\(\s*'([^']+)'\s*,\s*'([^']+)'\s*\)"
)

ATTACHMENT_FIELD_CANDIDATES_PATH = ["printFlpthNm", "flpthNm", "atchFileUrl"]
ATTACHMENT_FIELD_CANDIDATES_NAME = ["printFileNm", "fileNm", "atchFileNm"]
ATTACHMENT_RAW_HTML_CANDIDATES = ["fileDvNm", "atchFileHtml", "flieDown"]


def extract_attachment_url(item):
    """
    첨부파일 다운로드 URL과 원본 파일명을 반환. 없으면 (None, None).

    우선순위:
      1) onclick="fileBlank('path1' + 'path2' + 'path3', 'display_name.ext')"
         형태의 HTML이 필드 값에 그대로 들어있는 경우 정규식으로 파싱
      2) printFlpthNm(경로) + printFileNm(파일명) 필드 조합
    """
    # 1) HTML onclick 패턴이 필드 어딘가에 통째로 들어있는 경우
    for key in ATTACHMENT_RAW_HTML_CANDIDATES + ATTACHMENT_FIELD_CANDIDATES_PATH:
        val = item.get(key)
        if not val or not isinstance(val, str):
            continue
        if "fileBlank(" not in val:
            continue

        m = FILEBLANK_RE.search(val)
        if m:
            path = (m.group(1) + m.group(2) + m.group(3)).replace("//", "/")
            if not path.startswith("/"):
                path = "/" + path
            display_name = m.group(4)
            return f"https://{BIZINFO_HOST}{path}", display_name

        m2 = FILEBLANK_SIMPLE_RE.search(val)
        if m2:
            path = m2.group(1)
            if not path.startswith("/"):
                path = "/" + path
            display_name = m2.group(2)
            return f"https://{BIZINFO_HOST}{path}", display_name

    # 2) 경로 + 파일명 필드 조합 (일반적인 기업마당 API 구조)
    path_val = None
    for key in ATTACHMENT_FIELD_CANDIDATES_PATH:
        if item.get(key):
            path_val = str(item[key])
            break

    name_val = None
    for key in ATTACHMENT_FIELD_CANDIDATES_NAME:
        if item.get(key):
            name_val = str(item[key])
            break

    if path_val:
        # 이미 완전한 URL인 경우
        if path_val.startswith("http"):
            return path_val, (name_val or os.path.basename(path_val))
        # 상대 경로인 경우 도메인 결합, 중복 슬래시 정리
        full = f"https://{BIZINFO_HOST}/" + path_val.lstrip("/")
        full = re.sub(r"(?<!:)//+", "/", full.replace("https:/", "https://"))
        return full, (name_val or os.path.basename(path_val))

    return None, None


def download_attachment(url):
    """curl 우선, 실패 시 requests로 첨부파일 바이트 다운로드. (동일 도메인이므로 프록시도 동일 적용)"""
    try:
        cmd = ["curl", "-sSL", "--max-time", "30", "-A", "Mozilla/5.0 BizInfoBot/1.0"]
        if KR_PROXY_URL:
            cmd += ["--proxy", KR_PROXY_URL]
        cmd += [url]
        result = subprocess.run(cmd, capture_output=True, timeout=35)
        if result.returncode == 0 and result.stdout:
            return result.stdout
    except Exception as e:
        log(f"첨부파일 curl 다운로드 실패: {e}")

    try:
        proxies = {"https": KR_PROXY_URL, "http": KR_PROXY_URL} if KR_PROXY_URL else None
        resp = requests.get(
            url, timeout=30, headers={"User-Agent": "Mozilla/5.0"}, proxies=proxies
        )
        resp.raise_for_status()
        return resp.content
    except Exception as e:
        log(f"첨부파일 requests 다운로드 실패: {e}")
        return None


# =========================================================
# 4. Magic Bytes 판독 + 텍스트 추출
# =========================================================
def detect_file_type(raw_bytes):
    if not raw_bytes or len(raw_bytes) < 8:
        return "unknown"
    header = raw_bytes[:8]

    if header[:4] == b"%PDF":
        return "pdf"

    if header[:8] == b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1":
        return "hwp"  # 구버전 HWP (OLE Compound File)

    if header[:4] == b"PK\x03\x04":
        # zip 컨테이너 -> docx / hwpx 구분 필요
        try:
            with zipfile.ZipFile(io.BytesIO(raw_bytes)) as zf:
                names = zf.namelist()
                if any(n.startswith("word/") for n in names):
                    return "docx"
                if any("Contents/section" in n for n in names) or "mimetype" in names:
                    return "hwpx"
        except Exception:
            pass
        return "zip_unknown"

    return "unknown"


def extract_text_from_pdf(raw_bytes):
    if not pdfplumber:
        return ""
    text_parts = []
    tmp_path = os.path.join(TMP_DIR, f"tmp_{int(time.time()*1000)}.pdf")
    with open(tmp_path, "wb") as f:
        f.write(raw_bytes)
    try:
        try:
            with pdfplumber.open(tmp_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    text_parts.append(t)
        except Exception as e:
            # 손상되었거나 잘린 PDF(다운로드 실패 등)는 파싱 자체가 예외를
            # 던질 수 있음. 여기서 잡지 않으면 이 공고 하나 때문에 스크립트
            # 전체가 죽어서 그날 뉴스레터가 통째로 발송되지 않게 됨.
            log(f"PDF 파싱 실패(손상되었거나 잘린 파일일 수 있음): {e}")
            return ""

        text = "\n".join(text_parts).strip()

        # 텍스트가 거의 없으면(스캔본 PDF) OCR 시도
        if len(text) < 20 and pytesseract and convert_from_path:
            log("PDF 텍스트가 거의 없어 OCR을 시도합니다.")
            try:
                images = convert_from_path(tmp_path, dpi=200)
                ocr_texts = []
                for img in images[:5]:  # 과도한 처리 방지: 최대 5페이지
                    ocr_texts.append(pytesseract.image_to_string(img, lang="kor+eng"))
                text = "\n".join(ocr_texts).strip()
            except Exception as e:
                log(f"OCR 실패: {e}")
        return text
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass


def extract_text_from_docx(raw_bytes):
    if not docx:
        return ""
    try:
        f = io.BytesIO(raw_bytes)
        document = docx.Document(f)
        paragraphs = [p.text for p in document.paragraphs]
        # 표 안의 텍스트도 함께 추출
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs).strip()
    except Exception as e:
        log(f"DOCX 추출 실패: {e}")
        return ""


def extract_text_from_hwp(raw_bytes):
    """
    구버전 HWP(OLE) 파일에서 PrvText 스트림을 읽어 텍스트를 추출합니다.
    PrvText는 '미리보기 텍스트'이므로 전체 본문이 아닌 요약 수준일 수
    있습니다. 완전한 본문 추출이 필요하면 pyhwp(hwp5) 라이브러리 사용을
    권장합니다.
    """
    if not olefile:
        return ""
    try:
        f = io.BytesIO(raw_bytes)
        if not olefile.isOleFile(f):
            return ""
        ole = olefile.OleFileIO(f)
        if ole.exists("PrvText"):
            data = ole.openstream("PrvText").read()
            text = data.decode("utf-16-le", errors="ignore")
            return text.strip()
        return ""
    except Exception as e:
        log(f"HWP 추출 실패: {e}")
        return ""


def extract_text_from_hwpx(raw_bytes):
    """HWPX는 zip 기반 XML 포맷. section*.xml 내 텍스트 노드를 추출합니다."""
    try:
        text_parts = []
        with zipfile.ZipFile(io.BytesIO(raw_bytes)) as zf:
            section_files = sorted(
                [n for n in zf.namelist() if re.search(r"Contents/section\d*\.xml$", n)]
            )
            for name in section_files:
                xml_content = zf.read(name).decode("utf-8", errors="ignore")
                # <hp:t>...</hp:t> 또는 <t>...</t> 텍스트 노드 추출
                chunks = re.findall(r"<(?:hp:)?t[^>]*>(.*?)</(?:hp:)?t>", xml_content, re.DOTALL)
                cleaned = [re.sub(r"<[^>]+>", "", c) for c in chunks]
                text_parts.extend(cleaned)
        return "\n".join(text_parts).strip()
    except Exception as e:
        log(f"HWPX 추출 실패: {e}")
        return ""


def extract_text_from_attachment(raw_bytes):
    if not raw_bytes:
        return "", "unknown"
    file_type = detect_file_type(raw_bytes)
    log(f"첨부파일 타입 감지 결과: {file_type}")

    if file_type == "pdf":
        return extract_text_from_pdf(raw_bytes), file_type
    if file_type == "docx":
        return extract_text_from_docx(raw_bytes), file_type
    if file_type == "hwp":
        return extract_text_from_hwp(raw_bytes), file_type
    if file_type == "hwpx":
        return extract_text_from_hwpx(raw_bytes), file_type
    return "", file_type


# =========================================================
# 5. 이메일 / 전화번호 정규식 추출
# =========================================================
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"(0\d{1,2}[-.\s]?\d{3,4}[-.\s]?\d{4})")


# =========================================================
# 5-1. 경량 RAG: 관련 구절만 찾아서 넘기기 (NotebookLM 방식의 축소판)
# =========================================================
# 벡터DB/임베딩 없이도, 공고문 특성상 아래 키워드 주변만 잘라내면
# "지원규모/필수서류"에 필요한 내용을 대부분 커버할 수 있습니다.
# 이렇게 하면 (1) 프롬프트가 훨씬 작아지고 (2) 모델이 원문 전체를 훑으며
# reasoning 토큰을 낭비할 필요가 없어집니다.

# "지원규모"류 키워드: 금액/한도뿐 아니라, "최종 몇 개사 선정" 식의
# 선정 규모(정원)를 밝히는 공고도 많아 관련 표현을 폭넓게 포함시킴
BUDGET_KEYWORDS = [
    "지원규모", "지원 규모", "지원한도", "지원 한도", "지원금액", "지원 금액",
    "지원비율", "지원 비율", "보조금", "사업비", "지원한도액", "지원내용", "지원 내용",
    "사업규모", "사업 규모", "모집규모", "모집 규모", "선정규모", "선정 규모",
    "최종선정", "최종 선정", "선정인원", "선정 인원", "선정업체", "선정 업체",
    "모집인원", "모집 인원", "지원기업수", "지원 기업 수", "선정기업수",
]
# "필수서류"류 키워드
DOCS_KEYWORDS = [
    "제출서류", "제출 서류", "구비서류", "구비 서류", "필수서류", "필수 서류",
    "신청서류", "신청 서류", "첨부서류", "첨부 서류", "공통서류", "공통 서류",
    "제출 구비서류",
]

# 실제 공고문에 자주 등장하는 구체적 서류명 (직접 문자열 탐지용).
# 길이가 긴 표현을 먼저 검사해서, 짧은 표현이 긴 표현의 부분집합인 경우
# 중복으로 잡히지 않도록 처리합니다 (예: "사업자등록증명" vs "사업자등록증").
SPECIFIC_DOC_NAMES = [
    "사업자등록증명원", "사업자등록증명", "사업자등록증",
    "법인등기부등본", "법인 등기부등본",
    "표준재무제표증명", "표준재무제표증명원", "재무제표",
    "부가가치세과세표준증명원", "부가가치세과세표준증명",
    "국세완납증명서", "국세 완납증명서", "국세완납증명",
    "지방세완납증명서", "지방세 완납증명서", "지방세완납증명",
    "중소기업확인서",
]

# 지원 규모(금액/인원)를 나타내는 표현을 문서 전체에서 직접 정규식으로도 탐지
MONEY_RE = re.compile(
    r"(?:최대\s*)?\d[\d,]{0,12}\s*(?:천만원|백만원|만원|억원|원)"
    r"(?:\s*(?:이내|한도|내외))?"
)
SELECT_COUNT_RE = re.compile(
    r"\d+\s*(?:개사|개\s*기업|개\s*업체|개소|개\s*팀|명)\s*(?:내외)?"
    r"(?:을|를)?\s*(?:선정|모집|지원|선발)"
)


def _find_keyword_windows(text, keywords, before=30, after=350, max_per_keyword=2):
    """키워드가 등장하는 위치(키워드당 최대 max_per_keyword회) 주변을 구간으로 반환."""
    windows = []
    for kw in keywords:
        start_search = 0
        count = 0
        while count < max_per_keyword:
            idx = text.find(kw, start_search)
            if idx == -1:
                break
            start = max(0, idx - before)
            end = min(len(text), idx + len(kw) + after)
            windows.append((start, end))
            start_search = idx + len(kw)
            count += 1
    return windows


def _find_regex_windows(text, pattern, before=60, after=150, max_matches=3):
    """정규식 매치 위치 주변을 구간으로 반환 (금액/선정인원 등)."""
    windows = []
    for i, m in enumerate(pattern.finditer(text)):
        if i >= max_matches:
            break
        start = max(0, m.start() - before)
        end = min(len(text), m.end() + after)
        windows.append((start, end))
    return windows


def _merge_windows(windows, gap=50):
    """겹치거나 가까운 구간을 하나로 합쳐서 중복 텍스트를 줄임."""
    if not windows:
        return []
    windows.sort()
    merged = [windows[0]]
    for s, e in windows[1:]:
        last_s, last_e = merged[-1]
        if s <= last_e + gap:
            merged[-1] = (last_s, max(last_e, e))
        else:
            merged.append((s, e))
    return merged


def extract_relevant_context(text, max_chars=2600, fallback_chars=1500):
    """
    문서 전체 대신, [지원규모]/[필수서류] 관련 키워드·정규식 매치 주변
    구절만 뽑아서 반환합니다. 아무것도 못 찾으면 기존 방식(앞부분 일부)으로
    안전하게 폴백합니다.
    """
    if not text:
        return ""

    windows = (
        _find_keyword_windows(text, BUDGET_KEYWORDS)
        + _find_keyword_windows(text, DOCS_KEYWORDS)
        + _find_regex_windows(text, MONEY_RE)
        + _find_regex_windows(text, SELECT_COUNT_RE)
    )
    merged = _merge_windows(windows)

    if not merged:
        # 아무 단서도 못 찾음 -> 앞부분 일부로 폴백 (기존 동작 유지)
        return text[:fallback_chars]

    snippets = [text[s:e].strip() for s, e in merged]
    combined = "\n[...]\n".join(s for s in snippets if s)
    return combined[:max_chars]


def _dedupe_doc_names(found_names):
    """짧은 서류명이 이미 찾은 긴 서류명의 부분 문자열이면 제외."""
    found_sorted = sorted(set(found_names), key=len, reverse=True)
    deduped = []
    for name in found_sorted:
        if not any(name != other and name in other for other in deduped):
            deduped.append(name)
    return deduped


def deterministic_summary_hints(body_text):
    """
    정규식/문자열 매칭만으로 뽑아낸 결정론적 단서.
    - AI 호출이 실패해도 이 정보로 최소한의 요약을 만들 수 있고,
    - AI 호출이 성공할 때는 프롬프트에 '근거'로 함께 제공해 환각을 줄입니다.
    """
    if not body_text:
        return None, None

    money_matches = MONEY_RE.findall(body_text)[:3]
    select_matches = SELECT_COUNT_RE.findall(body_text)[:2]
    budget_hint = None
    hint_parts = []
    if money_matches:
        hint_parts.append(", ".join(dict.fromkeys(money_matches)))
    if select_matches:
        hint_parts.append(", ".join(dict.fromkeys(select_matches)))
    if hint_parts:
        budget_hint = " / ".join(hint_parts)

    found_docs = [name for name in SPECIFIC_DOC_NAMES if name in body_text]
    docs_hint = ", ".join(_dedupe_doc_names(found_docs)) if found_docs else None

    return budget_hint, docs_hint


# =========================================================
# 5-1-2. 원클릭 서비스 서류 카탈로그 매칭 (Y/N 플래그 + 기타서류)
# =========================================================
# "원클릭 서비스"가 자동 발급을 지원하는 서류 목록입니다. 공고문에서 이
# 서류명이 언급되면 Y, 아니면 N으로 엑셀에 표시하고, 매칭되지 않는
# 서류는 모두 "기타서류" 컬럼에 모아서 보여줍니다.
ONECLICK_DOC_CATALOG = [
    ("사업자등록증명", ["사업자등록증명원", "사업자등록증명"]),
    ("표준재무제표증명", ["표준재무제표증명원", "표준재무제표증명"]),
    ("부가가치세과세표준증명", ["부가가치세과세표준증명원", "부가가치세과세표준증명"]),
    ("납부내역증명", ["납부내역증명"]),
    ("납세증명", ["납세증명서", "납세증명"]),
    ("면세사업자수입금액증명원", ["면세사업자수입금액증명원", "면세사업자수입금액증명"]),
    ("사업자등록증명(영문)", ["사업자등록증명원(영문)", "영문사업자등록증명", "사업자등록증명(영문)", "영문 사업자등록증명"]),
    ("4대보험완납증명서", ["4대보험완납증명서", "4대 보험 완납증명서", "4대보험 완납증명"]),
    ("나의부동산정보", ["나의부동산정보", "부동산정보 통합열람"]),
    ("지방세 납세증명", ["지방세납세증명서", "지방세 납세증명"]),
    ("지방세 세목별 과세증명", ["지방세세목별과세증명서", "지방세 세목별 과세증명"]),
    ("법인 등기부등본", ["법인등기부등본", "법인 등기부등본"]),
    ("법인세 신고내역", ["법인세신고내역", "법인세 신고내역"]),
    ("종합소득세 신고내역", ["종합소득세신고내역", "종합소득세 신고내역"]),
    ("부가세 신고내역", ["부가세신고내역", "부가가치세신고내역", "부가세 신고내역"]),
    ("원천세 신고내역", ["원천세신고내역", "원천징수이행상황신고서", "원천세 신고내역"]),
    ("거래처별합계표", ["거래처별세금계산서합계표", "거래처별합계표"]),
    ("산업재해율확인서", ["산업재해율확인서", "산재보험료율확인서"]),
    ("산업재해 요양승인 및 반려여부 확인서", ["산업재해요양승인", "요양승인 및 반려여부", "산업재해 요양승인"]),
]

# "기타서류" 판별용으로, 문서 목록 후보를 뽑아낼 때 쓰는 제외 패턴
_ALL_CATALOG_KEYWORDS = [kw for _, kws in ONECLICK_DOC_CATALOG for kw in kws]


def compute_oneclick_doc_flags(body_text):
    """카탈로그의 각 서류가 본문에 언급되어 있는지 Y/N으로 판정."""
    flags = {}
    for label, keywords in ONECLICK_DOC_CATALOG:
        flags[label] = "Y" if (body_text and any(kw in body_text for kw in keywords)) else "N"
    return flags


def extract_document_list(text):
    """
    제출서류/구비서류/필수서류 등 섹션 근처에서 서류명 후보 목록을 추출합니다.
    번호/기호로 구분된 목록, 쉼표로 나열된 목록을 모두 최대한 지원하는
    휴리스틱이라 완벽하지는 않습니다.
    """
    if not text:
        return []
    windows = _find_keyword_windows(text, DOCS_KEYWORDS, before=0, after=500, max_per_keyword=3)
    merged = _merge_windows(windows)
    if not merged:
        return []
    section_text = "\n".join(text[s:e] for s, e in merged)

    # "제출서류:" 같은 섹션 헤더 문구 자체를 먼저 제거합니다. 이걸 안 하면
    # 헤더 문구와 실제 서류명이 한 덩어리로 붙어서(예: "제출서류: 사업계획서")
    # 아래 필터링 단계에서 통째로 걸러지는 문제가 생깁니다.
    header_pattern = "|".join(re.escape(k) for k in DOCS_KEYWORDS)
    section_text = re.sub(rf"(?:{header_pattern})\s*[:：]?", "\n", section_text)

    raw_parts = re.split(r"[\n,]", section_text)
    results = []
    for part in raw_parts:
        cleaned = re.sub(r"^[\s\-\・○□▪●\*0-9\.\)]+", "", part).strip()
        cleaned = re.sub(r"\s+", " ", cleaned)
        # 서류명치고 너무 짧거나 너무 긴 것, 한글이 없는 것은 제외
        if 2 <= len(cleaned) <= 30 and re.search(r"[가-힣]", cleaned):
            results.append(cleaned)
    return results


def compute_etc_docs(body_text, max_items=8):
    """카탈로그에 없는 서류명만 모아서 '기타서류' 값으로 사용할 문자열 생성."""
    doc_list = extract_document_list(body_text)
    etc = []
    for doc in doc_list:
        if any(kw in doc for kw in _ALL_CATALOG_KEYWORDS):
            continue
        etc.append(doc)
    etc_dedup = list(dict.fromkeys(etc))[:max_items]
    return ", ".join(etc_dedup) if etc_dedup else ""


# 한글 금액 단위를 실제 원(KRW) 숫자로 변환 (우선순위 산정용)
_WON_UNIT_MULTIPLIER = {"억": 100_000_000, "천만": 10_000_000, "백만": 1_000_000, "만": 10_000}
_WON_AMOUNT_RE = re.compile(r"(\d[\d,]*)\s*(억|천만|백만|만)?\s*원")


def parse_won_amount(text):
    """
    텍스트에서 발견되는 금액 표현 중 가장 큰 값을 원(KRW) 단위 숫자로
    반환합니다 (우선순위 계산용, 근사치입니다).
    """
    if not text:
        return 0
    best = 0
    for m in _WON_AMOUNT_RE.finditer(text):
        try:
            num = int(m.group(1).replace(",", ""))
        except ValueError:
            continue
        unit = m.group(2)
        multiplier = _WON_UNIT_MULTIPLIER.get(unit, 1)
        best = max(best, num * multiplier)
    return best


# 공고문 서식(신청서 양식)에는 "예) mail000@hanmail.net" 처럼 작성 예시용
# 플레이스홀더 이메일이 들어있는 경우가 많습니다. 이런 견본 주소를 실제 담당자
# 이메일로 잘못 뽑지 않도록 흔한 패턴을 걸러냅니다.
PLACEHOLDER_EMAIL_PATTERNS = [
    r"^mail\d+@",           # mail000@, mail123@ 등 서식 예시에서 자주 쓰이는 패턴
    r"^test\d*@",
    r"^sample\d*@",
    r"^example\d*@",
    r"^abc\d*@",
    r"^hong(gildong)?\d*@",  # '홍길동' 예시 이메일
    r"@example\.(com|co\.kr|net)$",
    r"@test\.(com|co\.kr|net)$",
]


def _is_placeholder_email(email):
    return any(re.search(p, email, re.IGNORECASE) for p in PLACEHOLDER_EMAIL_PATTERNS)


def extract_contacts(text):
    if not text:
        return [], []
    raw_emails = sorted(set(EMAIL_RE.findall(text)))
    emails = [e for e in raw_emails if not _is_placeholder_email(e)]
    phones = sorted(set(PHONE_RE.findall(text)))
    return emails, phones


# =========================================================
# 5-2. 유사(중복) 공고 통합
# =========================================================
# 기업마당에 동일/유사한 공고가 여러 건 올라오는 경우가 있어, 제목과
# 본문 내용이 모두 충분히 유사하면 하나로 합칩니다. 세부 프로그램명만
# 다른 진짜 별개의 공고(예: '맞춤형'/'마케팅'/'디자인' 역량강화)까지
# 잘못 합치지 않도록, 제목·본문 유사도를 모두 확인하는 이중 조건을
# 사용합니다. 임계값은 환경변수로 조정할 수 있습니다.
DEDUP_TITLE_THRESHOLD = float(os.environ.get("DEDUP_TITLE_THRESHOLD", "0.55"))
DEDUP_BODY_THRESHOLD = float(os.environ.get("DEDUP_BODY_THRESHOLD", "0.92"))


def _normalize_for_compare(text):
    if not text:
        return ""
    return re.sub(r"\s+", "", text)[:2000]


def _text_similarity(a, b):
    a_n, b_n = _normalize_for_compare(a), _normalize_for_compare(b)
    if not a_n or not b_n:
        return 0.0
    return difflib.SequenceMatcher(None, a_n, b_n).ratio()


def deduplicate_items(items_data):
    """
    title_sim(제목 유사도) AND body_sim(본문 유사도) 모두 임계값을 넘을 때만
    중복으로 판단해 통합합니다. 통합된 항목은 대표 항목의 '_dup_count'를
    늘리고 이유를 로그로 남깁니다.
    """
    kept = []
    for item in items_data:
        merged = False
        for k in kept:
            title_sim = _text_similarity(item["제목"], k["제목"])
            if title_sim < DEDUP_TITLE_THRESHOLD:
                continue
            body_sim = _text_similarity(
                item.get("_compare_text", ""), k.get("_compare_text", "")
            )
            if body_sim >= DEDUP_BODY_THRESHOLD:
                k["_dup_count"] = k.get("_dup_count", 1) + 1
                log(
                    "유사 공고로 판단되어 통합: "
                    f"'{item['제목']}' -> '{k['제목']}' "
                    f"(제목유사도={title_sim:.2f}, 본문유사도={body_sim:.2f})"
                )
                merged = True
                break
        if not merged:
            item["_dup_count"] = 1
            kept.append(item)
    return kept


# =========================================================
# 6. OpenRouter AI 요약
# =========================================================
def _clean_budget_output(raw_content):
    """
    모델 응답에서 지원규모를 나타내는 한 줄만 뽑아냅니다. 일부 무료 모델이
    "USER SAFETY: SAFE" 같은 안전 태그나 영어 메타 문구를 응답에 섞어
    보내는 경우가 있어, 그런 줄은 건너뛰고 실제 내용이 담긴 줄만 사용합니다.
    실패하면 None을 반환합니다.
    """
    if not raw_content:
        return None
    for line in raw_content.split("\n"):
        line = line.strip()
        if not line:
            continue
        # 순수 영문 대문자/기호로만 된 메타 문구성 줄은 건너뜀
        if re.fullmatch(r"[A-Z0-9\s:.,\-]+", line) and len(line) < 60:
            continue
        # 이모지/라벨 제거
        line = re.sub(r"^(📌\s*)?(지원\s*규모|사업\s*규모)\s*[:：]\s*", "", line)
        if line:
            return line
    return None


def summarize_budget(title, body_text):
    """
    원문에서 지원규모(지원금액 또는 선정 인원/기업 수)를 한국어 한 문장으로
    요약합니다. AI 호출이 불가능하거나 실패하면, 정규식으로 뽑아낸
    결정론적 단서(deterministic_summary_hints)를 그대로 반환합니다.
    """
    budget_hint, _ = deterministic_summary_hints(body_text)
    fallback = budget_hint if budget_hint else "확인 필요 (원문 참고)"

    if not OPENROUTER_API_KEY:
        return fallback

    if not body_text or len(body_text.strip()) < 10:
        return fallback

    # 원문 전체를 넘기는 대신, 지원규모 관련 구절만 검색해서 넘김 (경량 RAG)
    relevant_context = extract_relevant_context(body_text)
    if not relevant_context.strip():
        return fallback

    hint_block = (
        f"\n[참고용 자동 감지 단서 - 그대로 베끼지 말고 자연스럽게 반영]\n- {budget_hint}"
        if budget_hint else ""
    )

    prompt = (
        "다음은 정부/공공기관 지원사업 공고 원문에서 관련 부분만 발췌한 것입니다 "
        "(중간 생략은 [...]로 표시됨). 이 발췌문을 바탕으로 사업 지원규모(지원금액 "
        "또는 선정 인원/기업 수)를 한국어 한 문장으로만 답하세요. 다른 언어, 설명, "
        "태그, 메타 정보는 절대 포함하지 말고 이 한 문장만 출력하세요. 발췌문에 "
        "없는 내용은 추측하지 말고 '확인 필요'라고 쓰세요.\n\n"
        f"[공고 제목]\n{title}\n\n"
        f"[관련 발췌 부분]\n{relevant_context}"
        f"{hint_block}"
    )

    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": OPENROUTER_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        # 추론(reasoning) 토큰도 max_tokens 예산에서 차감되므로 충분히 크게 잡음.
        "max_tokens": 1000,
        "temperature": 0.2,
        "reasoning": {"max_tokens": 250, "exclude": True},
    }

    for attempt in range(1, 4):
        try:
            resp = requests.post(
                OPENROUTER_URL, headers=headers, json=payload, timeout=30
            )
            if resp.status_code != 200:
                log(
                    f"OpenRouter 요약 실패(시도 {attempt}): "
                    f"HTTP {resp.status_code} - {resp.text[:300]}"
                )
                resp.raise_for_status()

            data = resp.json()

            if "error" in data:
                log(f"OpenRouter 요약 실패(시도 {attempt}): API 에러 - {data['error']}")
                time.sleep(2 * attempt)
                continue

            choices = data.get("choices") or []
            if not choices:
                log(f"OpenRouter 요약 실패(시도 {attempt}): choices가 비어있음 - {str(data)[:300]}")
                time.sleep(2 * attempt)
                continue

            message = choices[0].get("message") or {}
            content = message.get("content") or message.get("reasoning")

            if not content or not str(content).strip():
                finish_reason = choices[0].get("finish_reason")
                log(
                    f"OpenRouter 요약 실패(시도 {attempt}): "
                    f"content가 비어있음 (finish_reason={finish_reason})"
                )
                time.sleep(2 * attempt)
                continue

            cleaned = _clean_budget_output(str(content))
            if cleaned:
                return cleaned

            log(
                f"OpenRouter 요약 실패(시도 {attempt}): "
                f"응답에서 유효한 줄을 찾지 못함 - 원본 일부: {str(content)[:200]!r}"
            )
            time.sleep(2 * attempt)
        except Exception as e:
            log(f"OpenRouter 요약 실패(시도 {attempt}): {e}")
            time.sleep(2 * attempt)

    return fallback


# =========================================================
# 7. HTML 뉴스레터 / 엑셀 생성
# =========================================================
# 엑셀/CSV 컬럼: 서류 카탈로그의 각 항목이 Y/N 플래그 컬럼으로 펼쳐짐
EXPORT_COLUMNS = (
    [
        ("공고명", "제목"),
        ("수행기관명(excInsttNm)", "수행기관명"),
        ("신청기간", "신청기간"),
        ("신청방법", "신청방법"),
        ("지원규모", "지원규모"),
    ]
    + [(label, label) for label, _ in ONECLICK_DOC_CATALOG]
    + [
        ("기타서류", "기타서류"),
        ("문의처", "문의처"),
        ("담당이메일", "담당 이메일"),
        ("담당전화번호", "담당 전화번호"),
        ("지원대상", "지원대상"),
        ("원문링크", "원문링크"),
    ]
)

# 컬럼별 엑셀 열 너비 (지정 안 된 컬럼은 기본값 사용)
_COLUMN_WIDTH_OVERRIDES = {
    "공고명": 40, "수행기관명(excInsttNm)": 22, "신청기간": 16, "신청방법": 26,
    "지원규모": 32, "기타서류": 32, "문의처": 22, "담당이메일": 24,
    "담당전화번호": 16, "지원대상": 20, "원문링크": 42,
}
_DEFAULT_COLUMN_WIDTH = 14


def build_xlsx_bytes(items_data):
    """
    전체 수집 데이터를 실제 엑셀(.xlsx) 파일 바이트로 만듭니다.
    (메일 본문에 data: URI로 넣지 않고 실제 첨부파일로 보내기 위해 raw bytes 반환)
    openpyxl이 없으면 None을 반환하며, 호출부에서 CSV로 폴백합니다.
    """
    if not openpyxl:
        return None

    headers = [label for label, _ in EXPORT_COLUMNS]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "기업마당 신규공고"
    ws.append(headers)

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="111827", end_color="111827", fill_type="solid")
    for col_idx in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(vertical="center", wrap_text=True)

    for row in items_data:
        ws.append([row.get(key, "") for _, key in EXPORT_COLUMNS])

    widths = [_COLUMN_WIDTH_OVERRIDES.get(label, _DEFAULT_COLUMN_WIDTH) for label, _ in EXPORT_COLUMNS]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    for row_cells in ws.iter_rows(min_row=2):
        for cell in row_cells:
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    ws.freeze_panes = "A2"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_csv_bytes(items_data):
    """전체 수집 데이터를 CSV 바이트로 만듭니다 (openpyxl 없을 때의 폴백용)."""
    output = io.StringIO()
    fieldnames = [label for label, _ in EXPORT_COLUMNS]
    writer = csv.writer(output)
    writer.writerow(fieldnames)
    for row in items_data:
        writer.writerow([row.get(key, "") for _, key in EXPORT_COLUMNS])
    return output.getvalue().encode("utf-8-sig")  # 엑셀 한글 깨짐 방지 BOM


def build_html_newsletter(items_data, target_date, attachment_filename=None):
    """
    PC 화면에서 보기 좋은 표(table) 형태의 정적 뉴스레터.
    컬럼: 기관명 / 공고명 / 사업규모 / 신청방법 / 원클릭매칭서류개수 /
          원클릭우선순위 / 공고바로가기
    - 모바일 대응은 고려하지 않고, PC 화면 폭을 그대로 활용합니다.
    - 원클릭 우선순위가 높은(숫자가 작은) 순서로 정렬해서 보여줍니다.
    """
    # 원클릭우선순위 오름차순(1순위가 위로) 정렬
    sorted_items = sorted(items_data, key=lambda r: r.get("_priority", 999999))

    rows_html = ""
    for i, row in enumerate(sorted_items):
        org_exc = row.get("수행기관명") or "-"
        title = row.get("제목") or "-"
        budget_display = (row.get("지원규모") or "-").split(" (")[0].strip()
        method = row.get("신청방법") or "-"
        match_count = row.get("_match_count", 0)
        priority = row.get("_priority", "-")
        link = row.get("원문링크") or "#"

        row_bg = "#ffffff" if i % 2 == 0 else "#f9fafb"

        rows_html += f"""
        <tr style="background:{row_bg};">
          <td style="padding:12px 14px;border-bottom:1px solid #e5e7eb;font-size:13.5px;color:#111827;white-space:nowrap;">{org_exc}</td>
          <td style="padding:12px 14px;border-bottom:1px solid #e5e7eb;font-size:13.5px;color:#111827;font-weight:600;">{title}</td>
          <td style="padding:12px 14px;border-bottom:1px solid #e5e7eb;font-size:13px;color:#374151;">{budget_display}</td>
          <td style="padding:12px 14px;border-bottom:1px solid #e5e7eb;font-size:13px;color:#374151;">{method}</td>
          <td style="padding:12px 14px;border-bottom:1px solid #e5e7eb;font-size:13.5px;color:#111827;text-align:center;">{match_count}건</td>
          <td style="padding:12px 14px;border-bottom:1px solid #e5e7eb;font-size:13.5px;color:#ffffff;text-align:center;">
            <span style="display:inline-block;min-width:22px;padding:3px 10px;background:#2563eb;border-radius:12px;font-weight:700;">{priority}</span>
          </td>
          <td style="padding:12px 14px;border-bottom:1px solid #e5e7eb;font-size:13px;text-align:center;">
            <a href="{link}" style="color:#2563eb;text-decoration:none;font-weight:600;">바로가기</a>
          </td>
        </tr>
        """

    table_html = f"""
    <table style="width:100%;border-collapse:collapse;background:#ffffff;
                  border-radius:10px;overflow:hidden;box-shadow:0 1px 3px rgba(0,0,0,0.06);">
      <thead>
        <tr style="background:#111827;">
          <th style="padding:12px 14px;color:#ffffff;font-size:13px;text-align:left;">기관명</th>
          <th style="padding:12px 14px;color:#ffffff;font-size:13px;text-align:left;">공고명</th>
          <th style="padding:12px 14px;color:#ffffff;font-size:13px;text-align:left;">사업규모</th>
          <th style="padding:12px 14px;color:#ffffff;font-size:13px;text-align:left;">신청방법</th>
          <th style="padding:12px 14px;color:#ffffff;font-size:13px;text-align:center;">원클릭매칭서류</th>
          <th style="padding:12px 14px;color:#ffffff;font-size:13px;text-align:center;">원클릭우선순위</th>
          <th style="padding:12px 14px;color:#ffffff;font-size:13px;text-align:center;">공고바로가기</th>
        </tr>
      </thead>
      <tbody>
        {rows_html if sorted_items else '<tr><td colspan="7" style="text-align:center;color:#6b7280;padding:40px 0;">오늘 등록된 신규 공고가 없습니다.</td></tr>'}
      </tbody>
    </table>
    """

    attachment_note = ""
    if attachment_filename:
        attachment_note = f"""
        <div style="text-align:center;margin-top:20px;padding:14px;background:#eef2ff;
                    border-radius:8px;font-size:13.5px;color:#3730a3;">
          📎 전체 상세 데이터(서류별 원클릭 발급 가능 여부 Y/N 포함)는
          첨부된 <b>{attachment_filename}</b> 파일에서 확인하실 수 있습니다.
        </div>
        """

    # PC 전용: 모바일 반응형 없이 넓은 고정 폭 사용 (표 컬럼이 7개라 여유 있게)
    html = f"""
    <html>
    <head><meta charset="utf-8"></head>
    <body style="margin:0;padding:0;background:#f3f4f6;font-family:'Apple SD Gothic Neo','Malgun Gothic',sans-serif;">
      <div style="max-width:1080px;margin:0 auto;padding:28px 24px;">
        <div style="text-align:center;margin-bottom:22px;">
          <h2 style="color:#111827;margin:0 0 6px 0;font-size:22px;">📢 기업마당 신규 지원사업 뉴스레터</h2>
          <div style="color:#6b7280;font-size:13.5px;">
            {target_date.strftime('%Y년 %m월 %d일')} 등록 공고 &middot; 총 {len(items_data)}건
            &middot; 원클릭매칭서류/우선순위 기준 정렬
          </div>
        </div>

        {table_html}

        {attachment_note}

        <div style="text-align:center;color:#9ca3af;font-size:11.5px;margin-top:24px;">
          본 메일은 기업마당(bizinfo.go.kr) 공개 API를 기반으로 자동 생성되었습니다.
        </div>
      </div>
    </body>
    </html>
    """
    return html


# =========================================================
# 8. Gmail SMTP 발송
# =========================================================
def parse_recipients(raw_receiver):
    """
    MAIL_RECEIVER 값을 안전하게 파싱합니다.
    - 콤마(,) 또는 세미콜론(;)으로 구분된 다중 수신자 지원
    - "홍길동 <hong@test.com>" 같은 이름 포함 형식도 지원
      (SMTP envelope에는 순수 이메일 주소만, 헤더 표시는 이름 포함 유지)
    - RFC 5321에 맞지 않는(순수 이메일 형식이 아닌) 항목은 걸러내고 경고 로그 출력
    반환값: (envelope_emails: list[str], header_value: str)
    """
    from email.utils import getaddresses, formataddr

    if not raw_receiver:
        return [], ""

    # 세미콜론을 콤마로 통일한 뒤, 빈 토큰(끝/중간의 연속 콤마 등)을 먼저
    # 제거합니다. getaddresses는 빈 토큰이 섞이면 전체 파싱이 깨질 수 있습니다.
    normalized = raw_receiver.replace(";", ",")
    tokens = [t.strip() for t in normalized.split(",") if t.strip()]
    parsed = getaddresses(tokens)

    valid = []
    for name, addr in parsed:
        addr = addr.strip()
        if EMAIL_RE.fullmatch(addr):
            valid.append((name.strip(), addr))
        else:
            log(f"경고: 유효하지 않은 수신자 형식이라 제외합니다 -> name='{name}' addr='{addr}'")

    envelope_emails = [addr for _, addr in valid]
    header_value = ", ".join(
        formataddr((name, addr)) if name else addr for name, addr in valid
    )
    return envelope_emails, header_value


def send_email(subject, html_body, attachment_bytes=None, attachment_filename=None):
    """
    attachment_bytes/attachment_filename이 주어지면 실제 첨부파일로 첨부합니다.
    (메일 본문의 data: URI 링크는 Gmail 등 대부분의 클라이언트가 보안상
    잘라내거나 무시하기 때문에, 다운로드가 필요한 파일은 반드시 진짜
    첨부파일로 보내야 합니다.)
    """
    if not (MAIL_USER and EMAIL_PASS and MAIL_RECEIVER):
        raise RuntimeError(
            "MAIL_USER / EMAIL_PASS / MAIL_RECEIVER 환경변수가 모두 필요합니다."
        )

    envelope_emails, header_value = parse_recipients(MAIL_RECEIVER)
    if not envelope_emails:
        raise RuntimeError(
            "MAIL_RECEIVER에서 유효한 이메일 주소를 하나도 찾지 못했습니다. "
            f"원본 값(참고용): {MAIL_RECEIVER!r}"
        )
    log(f"발송 대상 수신자 {len(envelope_emails)}명 확인: {envelope_emails}")

    msg = MIMEMultipart("mixed")
    msg["Subject"] = subject
    msg["From"] = MAIL_USER
    msg["To"] = header_value

    body_part = MIMEMultipart("alternative")
    body_part.attach(MIMEText(html_body, "html", "utf-8"))
    msg.attach(body_part)

    if attachment_bytes and attachment_filename:
        part = MIMEApplication(attachment_bytes, Name=attachment_filename)
        part["Content-Disposition"] = f'attachment; filename="{attachment_filename}"'
        msg.attach(part)
        log(f"첨부파일 추가됨: {attachment_filename} ({len(attachment_bytes)} bytes)")

    last_err = None
    for attempt in range(1, 4):
        try:
            with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=30) as server:
                server.login(MAIL_USER, EMAIL_PASS)
                # envelope 수신자는 반드시 순수 이메일 주소 리스트여야 함
                server.sendmail(MAIL_USER, envelope_emails, msg.as_string())
            log("메일 발송 성공.")
            return
        except Exception as e:
            last_err = e
            log(f"메일 발송 실패(시도 {attempt}): {e}")
            time.sleep(3 * attempt)

    raise RuntimeError(f"메일 발송 최종 실패: {last_err}")


# =========================================================
# 9. 메인 파이프라인
# =========================================================
def main():
    log("=== 기업마당 뉴스레터 봇 시작 ===")

    # 실행 시각(KST) 기준, "전날" 등록 공고를 대상으로 함
    kst = timezone(timedelta(hours=9))
    now_kst = datetime.now(kst)
    target_date = (now_kst - timedelta(days=1)).date()
    log(f"대상 등록일(전날): {target_date}")

    try:
        raw_items = fetch_bizinfo_items()
    except Exception as e:
        log(f"[치명적 오류] API 호출 실패: {e}")
        traceback.print_exc()
        sys.exit(1)

    new_items = filter_new_items(raw_items, target_date)

    items_data = []

    def process_single_item(item):
        """공고 1건을 처리해서 items_data용 row 딕셔너리를 반환. 실패 시 예외를 던짐."""
        title = item.get("pblancNm") or item.get("bsnsNm") or "(제목 없음)"
        org_exc = item.get("excInsttNm") or "-"
        period = item.get("reqstBeginEndDe") or item.get("pblancEndDe") or "-"
        link = item.get("pblancUrl") or item.get("rceptEngnHmpgUrl") or ""
        reqst_mth = strip_html(item.get("reqstMthPapersCn") or "") or "-"
        refrnc_nm = strip_html(item.get("refrncNm") or "")
        trget_nm = item.get("trgetNm") or "-"

        attachment_url, attachment_name = extract_attachment_url(item)
        body_text = ""
        if attachment_url:
            log(f"  첨부파일 발견: {attachment_name} ({attachment_url})")
            raw_bytes = download_attachment(attachment_url)
            body_text, file_type = extract_text_from_attachment(raw_bytes)
            if not body_text:
                log("  첨부파일 텍스트 추출 실패 또는 빈 결과.")
        else:
            # 첨부파일이 없으면 공고 요약 필드라도 사용
            body_text = item.get("bsnsSumryCn") or ""

        # 담당자 이메일/전화번호: API의 문의처(refrncNm) 필드를 우선 신뢰하고,
        # 거기서 못 찾으면 첨부파일 본문에서 정규식으로 추출 (서식 예시용
        # 플레이스홀더 이메일은 extract_contacts에서 자동으로 걸러짐)
        refrnc_emails, refrnc_phones = extract_contacts(refrnc_nm)
        body_emails, body_phones = extract_contacts(body_text)
        email = (refrnc_emails or body_emails or [""])[0]
        phone = (refrnc_phones or body_phones or [""])[0]

        # 지원규모(사업규모)만 AI로 한 줄 요약 (필수서류는 아래 결정론적
        # Y/N 플래그로 대체되어 더 이상 AI에게 요청하지 않음)
        budget = summarize_budget(title, body_text)

        # 원클릭 서비스가 지원하는 서류 카탈로그와 매칭해 Y/N 플래그 생성
        doc_flags = compute_oneclick_doc_flags(body_text)
        etc_docs = compute_etc_docs(body_text)
        match_count = sum(1 for v in doc_flags.values() if v == "Y")

        # 유사/중복 공고 판별에 쓸 비교용 텍스트 (제목 유사도와 함께 사용)
        compare_text = extract_relevant_context(body_text) or body_text[:1500]

        row = {
            "제목": title,
            "수행기관명": org_exc,
            "신청기간": period,
            "신청방법": reqst_mth,
            "지원규모": budget,
            "문의처": refrnc_nm or "-",
            "담당 이메일": email,
            "담당 전화번호": phone,
            "지원대상": trget_nm,
            "기타서류": etc_docs,
            "원문링크": link,
            "_compare_text": compare_text,
            "_match_count": match_count,
            "_budget_won": parse_won_amount(budget),
        }
        row.update(doc_flags)  # 카탈로그별 Y/N 컬럼 추가
        return row

    def fallback_row(item, error):
        """공고 처리 중 예외가 났을 때, 공고 자체는 누락시키지 않고
        최소 정보(제목/기관/원문링크)만 담은 대체 행을 만듭니다."""
        row = {
            "제목": item.get("pblancNm") or item.get("bsnsNm") or "(제목 없음)",
            "수행기관명": item.get("excInsttNm") or "-",
            "신청기간": item.get("reqstBeginEndDe") or item.get("pblancEndDe") or "-",
            "신청방법": "-",
            "지원규모": f"자동 처리 중 오류 발생 - 원문을 직접 확인해 주세요 ({error})",
            "문의처": "-",
            "담당 이메일": "",
            "담당 전화번호": "",
            "지원대상": "-",
            "기타서류": "",
            "원문링크": item.get("pblancUrl") or item.get("rceptEngnHmpgUrl") or "",
            "_compare_text": "",
            "_match_count": 0,
            "_budget_won": 0,
        }
        for label, _ in ONECLICK_DOC_CATALOG:
            row[label] = "N"
        return row

    for idx, item in enumerate(new_items, start=1):
        title_preview = item.get("pblancNm") or item.get("bsnsNm") or "(제목 없음)"
        log(f"[{idx}/{len(new_items)}] 처리 중: {title_preview}")

        # 공고 1건 처리 중 어떤 예외가 나더라도(손상된 첨부파일, 파싱 오류 등)
        # 이 공고만 건너뛰고 전체 배치는 계속 진행합니다. 여기서 잡지 않으면
        # 59건 중 1건만 실패해도 그날 뉴스레터 전체가 발송되지 않습니다.
        try:
            row = process_single_item(item)
        except Exception as e:
            log(f"  [경고] 이 공고 처리 중 예외 발생 - 건너뛰고 계속 진행합니다: {e}")
            traceback.print_exc()
            row = fallback_row(item, e)

        items_data.append(row)

    before_count = len(items_data)
    items_data = deduplicate_items(items_data)
    if len(items_data) != before_count:
        log(f"유사 공고 통합 결과: {before_count}건 -> {len(items_data)}건")

    # 원클릭우선순위 산정: 사업규모(금액) 순위 + 원클릭매칭서류개수 순위를
    # 합산해서 점수가 낮을수록(둘 다 상위일수록) 높은 우선순위를 부여합니다.
    # (근사치 산정 방식이며, 정확한 금액 비교가 필요하면 원문을 확인하세요.)
    if items_data:
        budget_order = sorted(
            range(len(items_data)), key=lambda i: -items_data[i]["_budget_won"]
        )
        match_order = sorted(
            range(len(items_data)), key=lambda i: -items_data[i]["_match_count"]
        )
        budget_rank = {i: r + 1 for r, i in enumerate(budget_order)}
        match_rank = {i: r + 1 for r, i in enumerate(match_order)}
        combined = sorted(
            range(len(items_data)),
            key=lambda i: budget_rank[i] + match_rank[i],
        )
        for priority, i in enumerate(combined, start=1):
            items_data[i]["_priority"] = priority

    # 엑셀(xlsx)을 실제 첨부파일로 준비 (본문 data: URI 링크는 Gmail 등에서
    # 잘리는 경우가 많아 신뢰할 수 없음 -> 반드시 진짜 첨부파일로 보냄)
    attachment_filename = f"bizinfo_{target_date.strftime('%Y%m%d')}.xlsx"
    attachment_bytes = build_xlsx_bytes(items_data)
    if not attachment_bytes:
        log("openpyxl 미설치 - CSV로 폴백합니다.")
        attachment_filename = f"bizinfo_{target_date.strftime('%Y%m%d')}.csv"
        attachment_bytes = build_csv_bytes(items_data)

    html_body = build_html_newsletter(
        items_data, target_date, attachment_filename=attachment_filename
    )
    subject = f"[기업마당 신규공고] {target_date.strftime('%Y-%m-%d')} 등록 {len(items_data)}건"

    try:
        send_email(
            subject, html_body,
            attachment_bytes=attachment_bytes,
            attachment_filename=attachment_filename,
        )
    except Exception as e:
        log(f"[치명적 오류] 메일 발송 실패: {e}")
        traceback.print_exc()
        sys.exit(1)

    log("=== 봇 실행 완료 ===")


if __name__ == "__main__":
    main()
